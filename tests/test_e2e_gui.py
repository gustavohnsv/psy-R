import os
import pytest
from PySide6.QtCore import Qt, QDate
from PySide6.QtWidgets import QMessageBox, QFileDialog

from main import MainWindow
from app.views import (
    TemplateScreen,
    PatientScreen,
    TemplateFieldsScreen,
    TestsScreen,
    ReviewScreen,
)

@pytest.fixture
def app(qtbot):
    """Fixture to create the application."""
    window = MainWindow()
    qtbot.addWidget(window)
    return window

def test_e2e_flow(app, qtbot, tmp_path, monkeypatch):
    """
    Test the full application flow:
    1. Load Template
    2. Fill Patient Data
    3. Fill Admin/Clinical/Behavior Data
    4. Fill Test Data
    5. Fill Conclusions
    6. Review and Generate Report
    """
    
    # Mock QFileDialog to return tmp_path
    monkeypatch.setattr(QFileDialog, "getExistingDirectory", lambda *args: str(tmp_path))
    
    # Mock QMessageBox to avoid blocking
    monkeypatch.setattr(QMessageBox, "information", lambda *args: QMessageBox.StandardButton.Ok)
    monkeypatch.setattr(QMessageBox, "warning", lambda *args: QMessageBox.StandardButton.Yes)
    monkeypatch.setattr(QMessageBox, "critical", lambda *args: QMessageBox.StandardButton.Ok)

    # 1. Template Screen
    assert isinstance(app.stacked_widget.currentWidget(), TemplateScreen)
    
    # Create a dummy template file
    template_path = tmp_path / "template.docx"
    from docx import Document
    doc = Document()
    doc.add_paragraph("Nome: {nome_paciente}")
    doc.add_paragraph("Idade: {idd_paciente}")
    doc.add_paragraph("Conclusão: {conclusao_text}")
    doc.save(template_path)
    
    # Simulate loading template
    app.template_screen.template_path = str(template_path)
    app.template_screen.file_template = doc
    # Manually trigger update since we bypassed the file dialog
    app.template_screen.ui.lineEdit_caminho_template.setText(str(template_path))
    app.template_screen._template_loaded = True
    app.template_screen.ui.btn_avancar.setEnabled(True)
    
    # Click Next
    qtbot.mouseClick(app.template_screen.ui.btn_avancar, Qt.LeftButton)
    
    # 2. Patient Screen
    assert isinstance(app.stacked_widget.currentWidget(), PatientScreen)
    app.patient_screen.ui.lineEdit_nome.setText("João Silva")
    app.patient_screen.ui.dateEdit_nascimento.setDate(QDate(2010, 1, 1))
    
    # Click Next
    qtbot.mouseClick(app.patient_screen.ui.btn_avancar, Qt.LeftButton)
    
    # 3. Admin Fields Screen
    assert isinstance(app.stacked_widget.currentWidget(), TemplateFieldsScreen)
    # Click Next
    qtbot.mouseClick(app.stacked_widget.currentWidget().ui.btn_avancar, Qt.LeftButton)
    
    # 4. Clinical Context Screen
    assert isinstance(app.stacked_widget.currentWidget(), TemplateFieldsScreen)
    # Click Next
    qtbot.mouseClick(app.stacked_widget.currentWidget().ui.btn_avancar, Qt.LeftButton)
    
    # 5. Behavior Screen
    assert isinstance(app.stacked_widget.currentWidget(), TemplateFieldsScreen)
    # Click Next
    qtbot.mouseClick(app.stacked_widget.currentWidget().ui.btn_avancar, Qt.LeftButton)
    
    # 6. Tests Screen
    assert isinstance(app.stacked_widget.currentWidget(), TestsScreen)
    # Click Next
    qtbot.mouseClick(app.tests_screen.ui.btn_avancar, Qt.LeftButton)
    
    # 7. Conclusions Section Screen
    assert isinstance(app.stacked_widget.currentWidget(), TemplateFieldsScreen)
    # Click Next (Navigate to Review)
    qtbot.mouseClick(app.stacked_widget.currentWidget().ui.btn_avancar, Qt.LeftButton)
    
    # 8. Review Screen
    assert isinstance(app.stacked_widget.currentWidget(), ReviewScreen)
    
    # Click Generate
    qtbot.mouseClick(app.review_screen.btn_generate_report, Qt.LeftButton)
    
    # Verify file creation
    output_file = tmp_path / "laudo_João_Silva.docx"
    assert output_file.exists()

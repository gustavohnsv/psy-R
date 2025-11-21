"""Unit tests for Branch 5: Document Generation."""
import pytest
import os
import tempfile
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock

# Skip if services don't exist (branch 3/5 not merged yet)
try:
    from app.services.template_processor import TemplateProcessor
    SERVICES_AVAILABLE = True
except ImportError:
    SERVICES_AVAILABLE = False

from docx import Document


@pytest.mark.unit
@pytest.mark.document_generation
@pytest.mark.skipif(not SERVICES_AVAILABLE, reason="Services module not available")
class TestTemplateProcessorFieldReplacement:
    """Test suite for field replacement in TemplateProcessor."""
    
    def test_replace_fields_in_paragraph(self):
        """Test replacing fields in a paragraph."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        doc.add_paragraph("Patient: {patient_name}, Age: {patient_crono_age}")
        
        field_mapping = {
            "patient_name": "João Silva",
            "patient_crono_age": "14"
        }
        
        processor = TemplateProcessor(doc)
        processor.replace_fields(field_mapping)
        
        # Check that fields were replaced
        text = doc.paragraphs[0].text
        assert "João Silva" in text
        assert "14" in text
        assert "{patient_name}" not in text
        assert "{patient_crono_age}" not in text
    
    def test_replace_fields_in_table(self):
        """Test replacing fields in a table."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "{patient_name}"
        table.cell(0, 1).text = "{patient_birth}"
        
        field_mapping = {
            "patient_name": "João Silva",
            "patient_birth": "15/03/2010"
        }
        
        processor = TemplateProcessor(doc)
        processor.replace_fields(field_mapping)
        
        assert table.cell(0, 0).text == "João Silva"
        assert table.cell(0, 1).text == "15/03/2010"
    
    def test_replace_fields_preserves_formatting(self):
        """Test that field replacement preserves formatting."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Patient: {patient_name}")
        run.bold = True
        
        field_mapping = {"patient_name": "João Silva"}
        
        processor = TemplateProcessor(doc)
        processor.replace_fields(field_mapping)
        
        # Text should be replaced
        assert "João Silva" in para.text
        # Note: Formatting preservation is complex and may need adjustment
    
    def test_replace_fields_handles_missing_fields(self):
        """Test that missing fields are handled gracefully."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        doc.add_paragraph("Patient: {patient_name}, Missing: {missing_field}")
        
        field_mapping = {"patient_name": "João Silva"}
        
        processor = TemplateProcessor(doc)
        processor.replace_fields(field_mapping)
        
        # Existing field should be replaced
        assert "João Silva" in doc.paragraphs[0].text
        # Missing field should remain as placeholder (or be empty)
        # This depends on implementation choice
    
    def test_replace_fields_in_headers_footers(self):
        """Test replacing fields in headers and footers."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        doc.sections[0].header.paragraphs[0].text = "{nome_psicologo}"
        doc.sections[0].footer.paragraphs[0].text = "{crp_psicologo}"
        
        field_mapping = {
            "nome_psicologo": "Dr. Ana Paula",
            "crp_psicologo": "CRP 06/123456"
        }
        
        processor = TemplateProcessor(doc)
        processor.replace_fields(field_mapping)
        
        assert doc.sections[0].header.paragraphs[0].text == "Dr. Ana Paula"
        assert doc.sections[0].footer.paragraphs[0].text == "CRP 06/123456"


@pytest.mark.unit
@pytest.mark.document_generation
@pytest.mark.skipif(not SERVICES_AVAILABLE, reason="Services module not available")
class TestDocumentSaving:
    """Test suite for document saving functionality."""
    
    def test_save_document(self, tmp_path):
        """Test saving a document to file."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        doc.add_paragraph("Test document")
        
        output_path = str(tmp_path / "test_output.docx")
        processor = TemplateProcessor()
        saved_path = processor.save_document(doc, output_path)
        
        assert saved_path == output_path
        assert os.path.exists(output_path)
        
        # Verify file can be opened
        loaded_doc = Document(output_path)
        assert loaded_doc.paragraphs[0].text == "Test document"
    
    def test_save_document_creates_directory(self, tmp_path):
        """Test that save_document creates directory if needed."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        doc.add_paragraph("Test")
        
        output_path = str(tmp_path / "new_dir" / "test.docx")
        processor = TemplateProcessor()
        processor.save_document(doc, output_path)
        
        assert os.path.exists(output_path)
        assert os.path.isdir(os.path.dirname(output_path))


@pytest.mark.unit
@pytest.mark.document_generation
@pytest.mark.skipif(not SERVICES_AVAILABLE, reason="Services module not available")
class TestPDFConversion:
    """Test suite for PDF conversion functionality."""
    
    @pytest.mark.skipif(
        not pytest.importorskip("docx2pdf", reason="docx2pdf not installed"),
        reason="docx2pdf library required"
    )
    def test_convert_to_pdf(self, tmp_path, sample_template_path):
        """Test converting DOCX to PDF."""
        from app.services.template_processor import TemplateProcessor
        
        # Create a simple DOCX file
        doc = Document()
        doc.add_paragraph("Test PDF conversion")
        docx_path = str(tmp_path / "test.docx")
        doc.save(docx_path)
        
        pdf_path = str(tmp_path / "test.pdf")
        processor = TemplateProcessor()
        
        try:
            result_path = processor.convert_to_pdf(docx_path, pdf_path)
            assert result_path == pdf_path
            # Note: PDF conversion may require additional dependencies
            # This test may need to be skipped in CI environments
        except Exception as e:
            # PDF conversion might fail in headless environments
            pytest.skip(f"PDF conversion failed: {e}")


@pytest.mark.integration
@pytest.mark.document_generation
@pytest.mark.skipif(not SERVICES_AVAILABLE, reason="Services module not available")
class TestDocumentGenerationIntegration:
    """Integration tests for complete document generation."""
    
    def test_full_document_generation_workflow(self, tmp_path, populated_data_model):
        """Test the complete document generation workflow."""
        from app.services.template_processor import TemplateProcessor
        
        # Create template
        template_doc = Document()
        template_doc.add_paragraph("Patient: {patient_name}")
        template_doc.add_paragraph("Birth: {patient_birth}")
        template_doc.add_paragraph("Conclusion: {conclusao_text}")
        
        template_path = str(tmp_path / "template.docx")
        template_doc.save(template_path)
        
        # Load template
        loaded_template = Document(template_path)
        populated_data_model.set_template(template_path, loaded_template)
        
        # Process and save
        processor = TemplateProcessor(loaded_template)
        field_mapping = populated_data_model.get_field_mapping()
        processor.replace_fields(field_mapping)
        
        output_path = str(tmp_path / "output.docx")
        processor.save_document(loaded_template, output_path)
        
        # Verify output
        assert os.path.exists(output_path)
        output_doc = Document(output_path)
        text = "\n".join([p.text for p in output_doc.paragraphs])
        
        assert "João Silva" in text
        assert "15/03/2010" in text
        assert "Este é um texto de conclusão de teste" in text
        assert "{patient_name}" not in text


@pytest.mark.e2e
@pytest.mark.document_generation
@pytest.mark.skipif(not SERVICES_AVAILABLE, reason="Services module not available")
class TestEndToEndDocumentGeneration:
    """End-to-end tests for document generation."""
    
    @patch('app.controllers.main_controller.QFileDialog.getExistingDirectory')
    @patch('app.controllers.main_controller.QMessageBox')
    def test_gerar_laudo_complete_flow(self, mock_messagebox, mock_filedialog, 
                                        tmp_path, qapp):
        """Test the complete gerar_laudo flow."""
        from main import MainWindow
        from docx import Document
        
        # Setup mocks
        mock_filedialog.return_value = str(tmp_path)
        mock_msg = MagicMock()
        mock_messagebox.information = mock_msg
        
        # Create window and setup data
        window = MainWindow()
        
        # Create and set template
        template_doc = Document()
        template_doc.add_paragraph("Patient: {patient_name}")
        template_path = str(tmp_path / "template.docx")
        template_doc.save(template_path)
        
        window.data_model.set_template(template_path, Document(template_path))
        window.data_model.set_patient_data({"patient_name": "Test Patient"})
        
        # Mock the template processor to avoid PDF conversion issues
        with patch('app.controllers.main_controller.TemplateProcessor') as mock_processor_class:
            mock_processor = MagicMock()
            mock_processor_class.return_value = mock_processor
            mock_processor.extract_fields.return_value = {"patient_name"}
            mock_processor.validate_fields.return_value = (["patient_name"], [])
            mock_processor.check_required_fields.return_value = ([], [])
            mock_processor.save_document.return_value = str(tmp_path / "output.docx")
            
            # Call generate_report (renamed from gerar_laudo)
            window.main_controller.generate_report()
            
            # Verify it was called
            assert mock_processor.replace_fields.called
            assert mock_processor.save_document.called

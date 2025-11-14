"""Unit tests for Branch 3: Template Field Extraction and Validation."""
import pytest
import sys
from pathlib import Path

# Skip if services don't exist (branch 3 not merged yet)
try:
    from app.services.field_validator import FieldValidator
    from app.services.template_processor import TemplateProcessor
    SERVICES_AVAILABLE = True
except ImportError:
    SERVICES_AVAILABLE = False

from docx import Document


@pytest.mark.unit
@pytest.mark.field_extraction
@pytest.mark.skipif(not SERVICES_AVAILABLE, reason="Services module not available")
class TestFieldValidator:
    """Test suite for FieldValidator class."""
    
    def test_validate_patient_fields(self):
        """Test validation of patient field names."""
        from app.services.field_validator import FieldValidator
        
        valid_fields = ["patient_name", "patient_birth", "patient_crono_age"]
        for field in valid_fields:
            is_valid, reason = FieldValidator.validate_field_name(field)
            assert is_valid, f"Field {field} should be valid: {reason}"
    
    def test_validate_respondent_fields(self):
        """Test validation of respondent field names."""
        from app.services.field_validator import FieldValidator
        
        valid_fields = ["resp1_name", "resp1_career", "resp2_name", "resp2_age"]
        for field in valid_fields:
            is_valid, reason = FieldValidator.validate_field_name(field)
            assert is_valid, f"Field {field} should be valid: {reason}"
    
    def test_validate_psychologist_fields(self):
        """Test validation of psychologist field names."""
        from app.services.field_validator import FieldValidator
        
        valid_fields = ["nome_psicologo", "crp_psicologo"]
        for field in valid_fields:
            is_valid, reason = FieldValidator.validate_field_name(field)
            assert is_valid, f"Field {field} should be valid: {reason}"
    
    def test_validate_test_fields(self):
        """Test validation of test field names."""
        from app.services.field_validator import FieldValidator
        
        valid_fields = ["QIT_out", "ICV_out", "F1_out", "ETDAH_CONCLUSAO_BLOCO"]
        for field in valid_fields:
            is_valid, reason = FieldValidator.validate_field_name(field)
            assert is_valid, f"Field {field} should be valid: {reason}"
    
    def test_reject_invalid_fields(self):
        """Test that truly invalid field names are rejected."""
        from app.services.field_validator import FieldValidator
        
        # Fields with invalid characters should be rejected
        invalid_fields = ["123field", "field-with-dash", "field.with.dot", "field with space", ""]
        for field in invalid_fields:
            is_valid, reason = FieldValidator.validate_field_name(field)
            assert not is_valid, f"Field '{field}' should be invalid: {reason}"
        
        # Valid format fields should be accepted (validator is permissive)
        valid_format_fields = ["invalid_field", "patient", "test_field", "random_name"]
        for field in valid_format_fields:
            is_valid, reason = FieldValidator.validate_field_name(field)
            assert is_valid, f"Field '{field}' has valid format and should be accepted: {reason}"
    
    def test_validate_fields_list(self):
        """Test validating a list of fields."""
        from app.services.field_validator import FieldValidator
        
        # All these fields have valid format (alphanumeric with underscores)
        fields = ["patient_name", "invalid_field", "resp1_name", "bad_field"]
        valid_fields, invalid_fields = FieldValidator.validate_fields(fields)
        
        # All should be valid since they follow valid format
        assert len(valid_fields) == 4
        assert "patient_name" in valid_fields
        assert "invalid_field" in valid_fields
        assert "resp1_name" in valid_fields
        assert "bad_field" in valid_fields
        assert len(invalid_fields) == 0
        
        # Test with truly invalid fields
        invalid_format_fields = ["123field", "field-with-dash", ""]
        valid_fields2, invalid_fields2 = FieldValidator.validate_fields(invalid_format_fields)
        assert len(invalid_fields2) > 0


@pytest.mark.unit
@pytest.mark.field_extraction
@pytest.mark.skipif(not SERVICES_AVAILABLE, reason="Services module not available")
class TestTemplateProcessor:
    """Test suite for TemplateProcessor class."""
    
    def test_extract_fields_from_paragraph(self):
        """Test extracting fields from a paragraph."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        doc.add_paragraph("Patient name: {patient_name}, Birth: {patient_birth}")
        
        processor = TemplateProcessor(doc)
        fields = processor.extract_fields()
        
        assert "patient_name" in fields
        assert "patient_birth" in fields
        assert len(fields) == 2
    
    def test_extract_fields_from_table(self):
        """Test extracting fields from a table."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "{patient_name}"
        table.cell(1, 0).text = "{patient_birth}"
        
        processor = TemplateProcessor(doc)
        fields = processor.extract_fields()
        
        assert "patient_name" in fields
        assert "patient_birth" in fields
    
    def test_extract_fields_from_headers_footers(self):
        """Test extracting fields from headers and footers."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        doc.sections[0].header.paragraphs[0].text = "{nome_psicologo}"
        doc.sections[0].footer.paragraphs[0].text = "{crp_psicologo}"
        
        processor = TemplateProcessor(doc)
        fields = processor.extract_fields()
        
        assert "nome_psicologo" in fields
        assert "crp_psicologo" in fields
    
    def test_extract_unique_fields(self):
        """Test that duplicate fields are only extracted once."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        doc.add_paragraph("{patient_name} and {patient_name} again")
        
        processor = TemplateProcessor(doc)
        fields = processor.extract_fields()
        
        assert "patient_name" in fields
        assert len(fields) == 1
    
    def test_validate_fields(self):
        """Test field validation."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        doc.add_paragraph("{patient_name} and {invalid_field}")
        
        processor = TemplateProcessor(doc)
        valid_fields, invalid_fields = processor.validate_fields()
        
        # Both fields have valid format, so both should be accepted
        assert "patient_name" in valid_fields
        assert "invalid_field" in valid_fields
        assert len(invalid_fields) == 0
        
        # Test with truly invalid field format
        doc2 = Document()
        doc2.add_paragraph("{patient_name} and {123invalid}")
        processor2 = TemplateProcessor(doc2)
        valid_fields2, invalid_fields2 = processor2.validate_fields()
        assert "patient_name" in valid_fields2
        assert len(invalid_fields2) > 0  # 123invalid should be rejected
    
    def test_check_required_fields(self):
        """Test checking for missing or empty fields."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        doc.add_paragraph("{patient_name} and {patient_birth}")
        
        processor = TemplateProcessor(doc)
        template_fields = processor.extract_fields()
        
        available_data = {"patient_name": "João", "patient_birth": ""}
        missing_fields, empty_fields = processor.check_required_fields(template_fields, available_data)
        
        assert len(missing_fields) == 0
        assert "patient_birth" in empty_fields
    
    def test_check_required_fields_missing(self):
        """Test checking for missing fields."""
        from app.services.template_processor import TemplateProcessor
        
        doc = Document()
        doc.add_paragraph("{patient_name} and {patient_birth}")
        
        processor = TemplateProcessor(doc)
        template_fields = processor.extract_fields()
        
        available_data = {"patient_name": "João"}
        missing_fields, empty_fields = processor.check_required_fields(template_fields, available_data)
        
        assert "patient_birth" in missing_fields
        assert len(empty_fields) == 0


@pytest.mark.integration
@pytest.mark.field_extraction
@pytest.mark.skipif(not SERVICES_AVAILABLE, reason="Services module not available")
class TestFieldExtractionIntegration:
    """Integration tests for field extraction."""
    
    def test_extract_and_validate_complex_template(self, tmp_path):
        """Test extracting and validating fields from a complex template."""
        from app.services.template_processor import TemplateProcessor
        
        # Create a complex template
        doc = Document()
        doc.add_paragraph("Patient: {patient_name}")
        doc.add_paragraph("Birth: {patient_birth}")
        
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Resp1: {resp1_name}"
        table.cell(1, 0).text = "Resp2: {resp2_name}"
        
        doc.sections[0].header.paragraphs[0].text = "Psychologist: {nome_psicologo}"
        
        processor = TemplateProcessor(doc)
        fields = processor.extract_fields()
        
        assert len(fields) >= 5
        assert "patient_name" in fields
        assert "patient_birth" in fields
        assert "resp1_name" in fields
        assert "resp2_name" in fields
        assert "nome_psicologo" in fields
        
        # Validate all fields
        valid_fields, invalid_fields = processor.validate_fields()
        assert len(invalid_fields) == 0  # All should be valid


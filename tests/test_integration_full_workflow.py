"""Integration tests for the complete application workflow."""
import pytest
import os
import tempfile
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from docx import Document
from PySide6.QtWidgets import QApplication
from PySide6.QtCore import QDate
import sys

# Ensure QApplication exists
if not QApplication.instance():
    app = QApplication(sys.argv)


@pytest.mark.integration
@pytest.mark.e2e
class TestCompleteWorkflow:
    """End-to-end tests for the complete application workflow."""
    
    def test_complete_workflow_from_template_to_document(self, tmp_path, qapp):
        """Test the complete workflow from loading template to generating document."""
        from main import MainWindow
        from app.models import ReportDataModel
        
        # Create a template file
        template_doc = Document()
        template_doc.add_paragraph("Laudo Psicológico")
        template_doc.add_paragraph("Paciente: {patient_name}")
        template_doc.add_paragraph("Data de Nascimento: {patient_birth}")
        template_doc.add_paragraph("Idade: {patient_crono_age}")
        template_doc.add_paragraph("Escola: {patient_school}")
        template_doc.add_paragraph("Turma: {patient_class}")
        template_doc.add_paragraph("")
        template_doc.add_paragraph("Responsável 1: {resp1_name}")
        template_doc.add_paragraph("Profissão: {resp1_career}")
        template_doc.add_paragraph("")
        template_doc.add_paragraph("Conclusão: {conclusion_text}")
        template_doc.add_paragraph("")
        template_doc.add_paragraph("Psicólogo: {psychologist_name}")
        template_doc.add_paragraph("CRP: {psychologist_crp}")
        
        template_path = str(tmp_path / "template.docx")
        template_doc.save(template_path)
        
        # Initialize application
        window = MainWindow()
        
        # Step 1: Load template

        window.template_screen.file_template = Document(template_path)
        window.template_screen._template_loaded = True
        window.template_screen.ui.lineEdit_caminho_template.setText(template_path)
        window.stacked_widget.setCurrentIndex(0)
        window.main_controller.collect_data_from_current_view()
        
        assert window.data_model.is_template_loaded()
        assert window.data_model.template_path == template_path
        
        # Step 2: Fill patient data
        window.stacked_widget.setCurrentIndex(1)
        window.patient_screen.ui.lineEdit_nome.setText("João Silva")
        window.patient_screen.ui.dateEdit_nascimento.setDate(QDate(2010, 3, 15))
        window.patient_screen.ui.lineEdit_idade_crono.setText("14")
        window.patient_screen.ui.lineEdit_escola.setText("Escola Municipal")
        window.patient_screen.ui.lineEdit_turma.setText("8º Ano")
        window.patient_screen.ui.lineEdit_resp1_nome.setText("Maria Silva")
        window.patient_screen.ui.lineEdit_resp1_profissao.setText("Professora")
        window.patient_screen.ui.lineEdit_resp1_escolaridade.setText("Superior")
        window.patient_screen.ui.lineEdit_resp1_idade.setText("35")
        window.main_controller.collect_data_from_current_view()
        
        assert window.data_model.patient_data["patient_name"] == "João Silva"
        assert window.data_model.resp1_data["resp1_name"] == "Maria Silva"
        
        # Step 3: Fill conclusion
        index_conclusao = window.stacked_widget.indexOf(window.conclusion_screen)
        window.stacked_widget.setCurrentIndex(index_conclusao)
        # Updated to access txt_conclusao directly as per new implementation
        window.conclusion_screen.txt_conclusao.setPlainText("Conclusão de teste")
        window.main_controller.collect_data_from_current_view()
        
        assert window.data_model.conclusion_text == "Conclusão de teste"
        
        # Step 4: Set psychologist data
        window.data_model.set_psychologist_data({
            "psychologist_name": "Dr. Ana Paula",
            "psychologist_crp": "CRP 06/123456"
        })
        
        # Step 5: Verify field mapping
        field_mapping = window.data_model.get_field_mapping()
        assert "patient_name" in field_mapping
        assert field_mapping["patient_name"] == "João Silva"
        assert "conclusion_text" in field_mapping
        assert "psychologist_name" in field_mapping
        
        # Step 6: Test review screen
        index_revisao = window.stacked_widget.indexOf(window.review_screen)
        window.stacked_widget.setCurrentIndex(index_revisao)
        window.review_screen.populate_summary()
        
        summary_text = window.review_screen.lbl_summary.text()
        assert "João Silva" in summary_text
        assert "Maria Silva" in summary_text
    
    @patch('app.controllers.main_controller.QFileDialog.getExistingDirectory')
    @patch('app.controllers.main_controller.QMessageBox')
    def test_document_generation_with_all_data(self, mock_messagebox, mock_filedialog,
                                                tmp_path, qapp):
        """Test generating document with all data filled."""
        from main import MainWindow
        from app.services.template_processor import TemplateProcessor
        
        # Setup
        mock_filedialog.return_value = str(tmp_path)
        
        # Create template
        template_doc = Document()
        template_doc.add_paragraph("{patient_name} - {patient_birth}")
        template_path = str(tmp_path / "template.docx")
        template_doc.save(template_path)
        
        # Setup window with data
        window = MainWindow()
        window.data_model.set_template(template_path, Document(template_path))
        window.data_model.set_patient_data({
            "patient_name": "Test Patient",
            "patient_birth": "01/01/2010"
        })
        window.data_model.set_conclusion_text("Test conclusion")
        
        # Mock processor to avoid PDF conversion
        with patch.object(TemplateProcessor, 'convert_to_pdf') as mock_pdf:
            mock_pdf.return_value = str(tmp_path / "output.pdf")
            
            # Generate document
            window.main_controller.generate_report()
            
            # Verify file was created
            docx_files = list(tmp_path.glob("*.docx"))
            assert len(docx_files) > 0


@pytest.mark.integration
class TestDataFlowBetweenScreens:
    """Test data flow between different screens."""
    
    def test_data_persistence_across_navigation(self, qapp):
        """Test that data persists when navigating between screens."""
        from main import MainWindow
        
        window = MainWindow()
        
        # Set data on patient screen
        window.stacked_widget.setCurrentIndex(1)
        window.patient_screen.ui.lineEdit_nome.setText("Test Patient")
        window.main_controller.collect_data_from_current_view()
        
        # Navigate away
        window.stacked_widget.setCurrentIndex(2)
        
        # Navigate back
        window.stacked_widget.setCurrentIndex(1)
        
        # Data should still be in model
        assert window.data_model.patient_data["patient_name"] == "Test Patient"
    
    def test_data_collection_on_each_screen(self, qapp):
        """Test that data is collected from each screen correctly."""
        from main import MainWindow
        
        window = MainWindow()
        
        # Template screen
        window.stacked_widget.setCurrentIndex(0)
        window.template_screen.file_template = Mock()
        window.template_screen._template_loaded = True
        window.template_screen.ui.lineEdit_caminho_template.setText("/test/path.docx")
        window.main_controller.collect_data_from_current_view()
        assert window.data_model.template_path == "/test/path.docx"
        
        # Patient screen
        window.stacked_widget.setCurrentIndex(1)
        window.patient_screen.ui.lineEdit_nome.setText("Patient")
        window.main_controller.collect_data_from_current_view()
        assert window.data_model.patient_data["patient_name"] == "Patient"
        
        # Conclusion screen
        index_conclusao = window.stacked_widget.indexOf(window.conclusion_screen)
        window.stacked_widget.setCurrentIndex(index_conclusao)
        window.conclusion_screen.txt_conclusao.setPlainText("Conclusion")
        window.main_controller.collect_data_from_current_view()
        assert window.data_model.conclusion_text == "Conclusion"

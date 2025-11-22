import pytest
from docx import Document

from app.models import ReportDataModel
from app.services.template_processor import TemplateProcessor


def import_template_loader():
    from app.services.template.loader import TemplateFieldsLoader

    return TemplateFieldsLoader


def test_template_fields_loader_includes_expected_fields():
    TemplateFieldsLoader = import_template_loader()
    loader = TemplateFieldsLoader()

    config = loader.load_config()
    assert "sections" in config and config["sections"], "Config must list sections"

    fields = loader.get_all_fields()
    assert "analise_paciente" in fields

    analise_field = fields["analise_paciente"]
    assert analise_field["widget"] == "textarea"

    assert all(
        not any(
            token in name
            for token in ("WISC", "RAVLT", "ETDAH", "SRS", "CARS", "HTP", "BPA")
        )
        for name in fields
    ), "Test-specific fields should not be present in the config"


def test_data_model_maps_template_fields_into_mapping():
    data_model = ReportDataModel()
    payload = {
        "analise_paciente": "Observações relevantes",
        "cidade": "Poá",
    }

    data_model.set_template_field_values(payload)
    mapping = data_model.get_field_mapping()

    for key, value in payload.items():
        assert mapping[key] == value


def test_template_processor_replaces_custom_fields_from_mapping():
    doc = Document()
    doc.add_paragraph("Resumo: {analise_paciente}")
    doc.add_paragraph("Local: {cidade}")

    processor = TemplateProcessor(doc)

    data_model = ReportDataModel()
    data_model.set_template_field_values(
        {"analise_paciente": "Paciente atento", "cidade": "Poá"}
    )

    processor.replace_fields(data_model.get_field_mapping(), doc)

    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "Paciente atento" in text
    assert "Poá" in text


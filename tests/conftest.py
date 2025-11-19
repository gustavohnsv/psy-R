"""Pytest configuration and shared fixtures."""
import sys
import os
from pathlib import Path
import pytest
from unittest.mock import Mock, MagicMock
from docx import Document

# Add src to path
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

@pytest.fixture
def sample_template_path(tmp_path):
    """Create a sample DOCX template file for testing."""
    template_path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("Template with {patient_name} and {patient_birth}")
    doc.save(str(template_path))
    return str(template_path)

@pytest.fixture
def sample_document():
    """Create a sample Document object."""
    doc = Document()
    doc.add_paragraph("Test document with {patient_name}")
    return doc

@pytest.fixture
def sample_patient_data():
    """Sample patient data for testing."""
    return {
        "patient_name": "João Silva",
        "patient_birth": "15/03/2010",
        "patient_crono_age": "14",
        "patient_school": "Escola Municipal",
        "patient_class": "8º Ano"
    }

@pytest.fixture
def sample_resp1_data():
    """Sample respondent 1 data for testing."""
    return {
        "resp1_name": "Maria Silva",
        "resp1_career": "Professora",
        "resp1_education": "Superior Completo",
        "resp1_age": 35
    }

@pytest.fixture
def sample_resp2_data():
    """Sample respondent 2 data for testing."""
    return {
        "resp2_name": "José Silva",
        "resp2_career": "Engenheiro",
        "resp2_education": "Superior Completo",
        "resp2_age": 38
    }

@pytest.fixture
def sample_test_results():
    """Sample test results for testing."""
    return {
        "QIT_WISC": 116,
        "ICV_WISC": 110,
        "IOP_WISC": 100,
        "F1_ETDAH": 82,
        "SRS_ESCORE_TOTAL": 61,
        "CARS_PONTUACAO": 30,
        "TASK_NEUP": 72,
        "ETDAH_CONCLUSAO_BLOCO": "Resultado esperado"
    }

@pytest.fixture
def sample_conclusion_text():
    """Sample conclusion text for testing."""
    return "Este é um texto de conclusão de teste."

@pytest.fixture
def sample_psychologist_data():
    """Sample psychologist data for testing."""
    return {
        "nome_psicologo": "Dr. Ana Paula",
        "crp_psicologo": "CRP 06/123456"
    }

@pytest.fixture
def populated_data_model(sample_patient_data, sample_resp1_data, sample_resp2_data, 
                         sample_test_results, sample_conclusion_text, sample_psychologist_data,
                         sample_document):
    """Create a fully populated data model for testing."""
    from app.models import LaudoDataModel
    
    model = LaudoDataModel()
    model.set_template("/test/path/template.docx", sample_document)
    model.set_patient_data(sample_patient_data)
    model.set_resp1_data(sample_resp1_data)
    model.set_resp2_data(sample_resp2_data)
    model.set_test_results(sample_test_results)
    model.set_conclusion_text(sample_conclusion_text)
    model.set_psychologist_data(sample_psychologist_data)
    
    return model

@pytest.fixture
def qapp(qtbot):
    """Create QApplication for GUI tests."""
    from PySide6.QtWidgets import QApplication
    import sys
    
    if not QApplication.instance():
        app = QApplication(sys.argv)
    else:
        app = QApplication.instance()
    
    return app


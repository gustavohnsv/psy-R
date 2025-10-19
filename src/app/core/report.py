import json
from typing import Dict
from .models import ReportModel
from statistics import mean
import math


def convert_scale(test_id: str, raw: Dict) -> Dict:
    """
    Placeholder: aplicar regras de conversão por teste.
    Troque por regras reais das escalas.
    """
    converted = {}
    for k, v in raw.items():
        try:
            converted[k] = float(v)
        except Exception:
            converted[k] = v

    # Exemplo de regra simples
    if test_id == "exemplo_int":
        converted["score_t"] = converted.get("raw_score", 0) + 50
    return converted


def fill_report_model(model: ReportModel) -> ReportModel:
    for t in model.tests:
        t.converted_values = convert_scale(t.test_id, t.raw_values)
    return model


def export_to_json(model: ReportModel, path: str):
    with open(path, "w", encoding="utf-8") as f:
        json.dump({
            "patient": model.patient.__dict__,
            "tests": [
                {"test_id": t.test_id, "raw": t.raw_values, "converted": t.converted_values}
                for t in model.tests
            ],
            "metadata": model.metadata,
        }, f, ensure_ascii=False, indent=2)


def export_to_docx(model: ReportModel, path: str):
    """Exporta um relatório simples para DOCX usando placeholders.
    Se python-docx não estiver instalado, levanta ImportError.
    """
    try:
        from docx import Document
    except Exception as e:
        raise ImportError("python-docx is required for DOCX export") from e

    doc = Document()
    doc.add_heading('Laudo - Resumo', level=1)
    p = doc.add_paragraph()
    p.add_run('Paciente: ').bold = True
    p.add_run(model.patient.name)

    doc.add_heading('Dados do Paciente', level=2)
    doc.add_paragraph(f"Idade: {model.patient.age}")
    doc.add_paragraph(f"Data de nascimento: {model.patient.birthdate}")
    doc.add_paragraph(f"Escola: {model.patient.school}")
    doc.add_paragraph(f"Turma/Grupo: {model.patient.group}")

    doc.add_heading('Testes', level=2)
    for t in model.tests:
        doc.add_heading(t.test_id, level=3)
        doc.add_paragraph('Valores brutos:')
        for k, v in t.raw_values.items():
            doc.add_paragraph(f"- {k}: {v}")
        doc.add_paragraph('Valores convertidos:')
        for k, v in t.converted_values.items():
            doc.add_paragraph(f"- {k}: {v}")

    doc.save(path)


def compute_mean_and_percentiles(values_list):
    """Recebe uma lista de números e retorna média e percentis 25/50/75/90.
    Valores inválidos são ignorados.
    Retorna dict {"mean": ..., "p25": ..., "p50": ..., "p75": ..., "p90": ...}
    """
    nums = []
    for v in values_list:
        try:
            nums.append(float(v))
        except Exception:
            continue
    if not nums:
        return {"mean": None, "p25": None, "p50": None, "p75": None, "p90": None}
    nums.sort()
    def percentile(arr, p):
        if not arr:
            return None
        k = (len(arr) - 1) * (p / 100.0)
        f = math.floor(k)
        c = math.ceil(k)
        if f == c:
            return arr[int(k)]
        d0 = arr[int(f)] * (c - k)
        d1 = arr[int(c)] * (k - f)
        return d0 + d1

    return {
        "mean": mean(nums),
        "p25": percentile(nums, 25),
        "p50": percentile(nums, 50),
        "p75": percentile(nums, 75),
        "p90": percentile(nums, 90),
    }


def fill_docx_template(template_path: str, model: ReportModel, out_path: str):
    """Carrega um template DOCX com placeholders no formato {{campo}} e substitui
    por valores do modelo. Gera um novo DOCX em out_path.
    Implementação simples: substitui runs de texto que contenham {{placeholder}}.
    """
    try:
        from docx import Document
    except Exception as e:
        raise ImportError("python-docx is required for template filling") from e

    doc = Document(template_path)

    # mapa simples de placeholders
    placeholders = {
        "patient_name": model.patient.name,
        "patient_age": model.patient.age,
        "patient_birthdate": model.patient.birthdate,
        "patient_school": model.patient.school or "",
        "patient_group": model.patient.group or "",
    }

    # adicionar resumo de testes como texto único
    tests_summary = []
    for t in model.tests:
        tests_summary.append(f"{t.test_id}: {t.converted_values}")
    placeholders["tests_summary"] = "\n".join(tests_summary)

    for para in doc.paragraphs:
        for key, val in placeholders.items():
            token = f"{{{{{key}}}}}"
            if token in para.text:
                # substituir nos runs para preservar formatação
                inline = para.runs
                for i in range(len(inline)):
                    if token in inline[i].text:
                        inline[i].text = inline[i].text.replace(token, str(val))

    # também processar tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, val in placeholders.items():
                        token = f"{{{{{key}}}}}"
                        if token in para.text:
                            for run in para.runs:
                                if token in run.text:
                                    run.text = run.text.replace(token, str(val))

    doc.save(out_path)


def export_to_pdf(model: ReportModel, path: str):
    """Tenta gerar um PDF a partir de um DOCX preenchido.
    Primeiro tenta docx2pdf (Windows). Se não disponível, gera um PDF simples com reportlab.
    """
    try:
        # criar um DOCX temporário e converter
        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        tmp.close()
        fill_docx_template(model.metadata.get('template_path', ''), model, tmp.name)
        try:
            from docx2pdf import convert
            convert(tmp.name, path)
            return
        except Exception:
            pass

        # fallback: gerar um PDF texto simples
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
        except Exception:
            raise RuntimeError('Nenhuma biblioteca disponível para gerar PDF (docx2pdf/reportlab)')

        c = canvas.Canvas(path, pagesize=A4)
        w, h = A4
        y = h - 40
        c.setFont('Helvetica-Bold', 14)
        c.drawString(40, y, 'Laudo psicológico - Resumo')
        y -= 30
        c.setFont('Helvetica', 10)
        lines = []
        lines.append(f"Paciente: {model.patient.name}")
        lines.append(f"Idade: {model.patient.age}")
        lines.append(f"Nascimento: {model.patient.birthdate}")
        lines.append('')
        for t in model.tests:
            lines.append(f"{t.test_id}: {t.converted_values}")

        for ln in lines:
            if y < 60:
                c.showPage()
                y = h - 40
            c.drawString(40, y, str(ln))
            y -= 14
        c.save()
    finally:
        try:
            import os
            os.unlink(tmp.name)
        except Exception:
            pass


def export_to_txt(model: ReportModel, path: str):
    """Exporta um relatório simples para texto (.txt) — fácil de abrir no Word."""
    lines = []
    lines.append("LAUDO PSICOLÓGICO - RESUMO")
    lines.append("=" * 40)
    lines.append(f"Paciente: {model.patient.name}")
    lines.append(f"Idade: {model.patient.age}")
    lines.append(f"Data de nascimento: {model.patient.birthdate}")
    lines.append(f"Escola: {model.patient.school}")
    lines.append(f"Turma/Grupo: {model.patient.group}")
    lines.append("")
    lines.append("TESTES")
    lines.append("-" * 40)
    for t in model.tests:
        lines.append(f"Teste: {t.test_id}")
        lines.append("  Valores brutos:")
        if t.raw_values:
            for k, v in t.raw_values.items():
                lines.append(f"    {k}: {v}")
        else:
            lines.append("    (nenhum)")
        lines.append("  Valores convertidos:")
        if t.converted_values:
            for k, v in t.converted_values.items():
                lines.append(f"    {k}: {v}")
        else:
            lines.append("    (nenhum)")
        lines.append("-" * 40)

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

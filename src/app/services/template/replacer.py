import re
from typing import Dict, List, Tuple
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

class TemplateReplacer:
    """Service responsible for replacing fields in DOCX templates."""
    
    FIELD_PATTERN = re.compile(r'\{([a-zA-Z0-9_]+)\}')
    
    def replace_fields(self, document: Document, field_mapping: Dict[str, str]) -> int:
        """Replace all field placeholders in the document with actual values.
        
        Preserves formatting by handling fields that may be split across multiple runs.
        
        Returns:
            Total number of replacements made
        """
        if document is None:
            raise ValueError("No document provided for field replacement")
        
        replacement_count = 0
        
        # Replace in main document paragraphs
        for paragraph in document.paragraphs:
            count = self._replace_in_paragraph(paragraph, field_mapping)
            replacement_count += count
        
        # Replace in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        count = self._replace_in_paragraph(paragraph, field_mapping)
                        replacement_count += count
        
        # Replace in headers
        for section in document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                count = self._replace_in_paragraph(paragraph, field_mapping)
                replacement_count += count
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            count = self._replace_in_paragraph(paragraph, field_mapping)
                            replacement_count += count
        
        # Replace in footers
        for section in document.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                count = self._replace_in_paragraph(paragraph, field_mapping)
                replacement_count += count
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            count = self._replace_in_paragraph(paragraph, field_mapping)
                            replacement_count += count
        
        return replacement_count

    def _replace_in_paragraph(self, paragraph: Paragraph, field_mapping: Dict[str, str]) -> int:
        """Replace fields in a paragraph, handling fields split across runs."""
        full_text = paragraph.text
        matches = list(self.FIELD_PATTERN.finditer(full_text))
        
        if not matches:
            return 0
        
        replacement_count = 0
        
        # Work backwards to preserve indices
        for match in reversed(matches):
            field_name = match.group(1)
            start_pos = match.start()
            end_pos = match.end()
            
            if field_name in field_mapping:
                replacement = str(field_mapping[field_name]) if field_mapping[field_name] else ""
                self._replace_field_in_runs(paragraph, start_pos, end_pos, replacement)
                replacement_count += 1
            else:
                # Field not in mapping - replace with empty string
                self._replace_field_in_runs(paragraph, start_pos, end_pos, "")
                replacement_count += 1
        
        return replacement_count
    
    def _replace_field_in_runs(self, paragraph: Paragraph, start_pos: int, end_pos: int, replacement: str):
        """Replace a field that may span multiple runs in a paragraph."""
        runs_data = []
        current_pos = 0
        
        for run in paragraph.runs:
            run_text = run.text
            run_start = current_pos
            run_end = current_pos + len(run_text)
            runs_data.append((run, run_start, run_end, run_text))
            current_pos = run_end
        
        target_runs = []
        for run, run_start, run_end, run_text in runs_data:
            if not (run_end <= start_pos or run_start >= end_pos):
                target_runs.append((run, run_start, run_end, run_text))
        
        if not target_runs:
            return
        
        if len(target_runs) == 1:
            run, run_start, run_end, run_text = target_runs[0]
            rel_start = max(0, start_pos - run_start)
            rel_end = min(len(run_text), end_pos - run_start)
            new_text = run_text[:rel_start] + replacement + run_text[rel_end:]
            run.text = new_text
        else:
            first_run, first_start, first_end, first_text = target_runs[0]
            rel_start = max(0, start_pos - first_start)
            first_run.text = first_text[:rel_start] + replacement
            
            for run, _, _, _ in target_runs[1:-1]:
                run.text = ""
            
            if len(target_runs) > 1:
                last_run, last_start, last_end, last_text = target_runs[-1]
                rel_end = min(len(last_text), end_pos - last_start)
                last_run.text = last_text[rel_end:]

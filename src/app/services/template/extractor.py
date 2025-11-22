import re
from typing import Set, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

class TemplateExtractor:
    """Service responsible for extracting fields from DOCX templates."""
    
    FIELD_PATTERN = re.compile(r'\{([a-zA-Z0-9_]+)\}')
    
    def extract_fields(self, document: Document) -> Set[str]:
        """Extract all field names from a DOCX document.
        
        Searches through paragraphs, tables, headers, and footers.
        """
        if document is None:
            return set()
        
        fields = set()
        
        # Extract from main document paragraphs
        for paragraph in document.paragraphs:
            fields.update(self._extract_from_paragraph(paragraph))
        
        # Extract from tables
        for table in document.tables:
            fields.update(self._extract_from_table(table))
        
        # Extract from headers
        for section in document.sections:
            header = section.header
            for paragraph in header.paragraphs:
                fields.update(self._extract_from_paragraph(paragraph))
            
            # Extract from header tables
            for table in header.tables:
                fields.update(self._extract_from_table(table))
        
        # Extract from footers
        for section in document.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                fields.update(self._extract_from_paragraph(paragraph))
            
            # Extract from footer tables
            for table in footer.tables:
                fields.update(self._extract_from_table(table))
        
        return fields
    
    def _extract_from_paragraph(self, paragraph: Paragraph) -> Set[str]:
        """Extract field names from a paragraph."""
        fields = set()
        text = paragraph.text
        
        # Find all field patterns in the paragraph text
        matches = self.FIELD_PATTERN.findall(text)
        fields.update(matches)
        
        return fields
    
    def _extract_from_table(self, table: Table) -> Set[str]:
        """Extract field names from a table."""
        fields = set()
        
        for row in table.rows:
            for cell in row.cells:
                # Extract from paragraphs in each cell
                for paragraph in cell.paragraphs:
                    fields.update(self._extract_from_paragraph(paragraph))
        
        return fields
